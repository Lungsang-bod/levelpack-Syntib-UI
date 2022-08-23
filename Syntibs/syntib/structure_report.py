from collections import defaultdict

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from nltk import ParentedTree, Tree


def get_structures(tree, structures, total_structures, tree_name):
    tree = cleanup(tree)
    ptree = ParentedTree.convert(tree)

    sentence = ''.join(tree.leaves()) + f' — sent. {tree_name}'
    if sentence not in structures:
        structures[sentence] = {}
    if sentence not in total_structures:
        total_structures[sentence] = {}

    filter_leaves = lambda t: t.height() >= 3
    for subtree in ptree.subtrees(filter_leaves):
        label = subtree.label()
        if label not in structures[sentence]:
            structures[sentence][label] = {}
        if label not in total_structures[sentence]:
            total_structures[sentence][label] = {}

        # label: label + labels of direct children
        sub_label = f'{subtree.label()}({" ".join([s.label() for s in subtree])})'
        if sub_label not in structures[sentence][label]:
            structures[sentence][label][sub_label] = []
        if sub_label not in total_structures[sentence][label]:
            total_structures[sentence][label][sub_label] = []

        # sub
        sub_str = ''.join(subtree.leaves())
        sub_struct = f'   [{subtree.label()} ' + ' '.join([f'{sub.flatten()}' for sub in subtree]) + ']   '
        location = sentence.replace(sub_str, sub_struct).strip()

        structures[sentence][label][sub_label].append((location, str(subtree), sub_struct.strip()))
        total_structures[sentence][label][sub_label].append((location, str(subtree), sub_struct.strip()))


def cleanup(tree):
    def remove_actancy(tree):
        if '-' in tree.label():
            tree.set_label(tree.label().split('-')[-1])
        for sub in tree:
            if isinstance(tree, Tree) and tree.height() >= 3:
                remove_actancy(sub)

    def remove_punct(tree):
        filtered_children = []
        for t in tree:
            if t.label() != 'PUNCT':
                filtered_children.append(t)
        tree = Tree(tree.label(), filtered_children)

        for sub in tree:
            if isinstance(sub, Tree) and tree.height() >= 3:
                remove_actancy(sub)

        return tree

    tree = tree.copy(deep=True)
    remove_actancy(tree)
    tree = remove_punct(tree)
    return tree


def gen_struct_sentence_report(structures, out_file):
    def gen_summary(a):
        total_s = 0
        total_labels = defaultdict(int)
        total_structs = defaultdict(int)
        for label, b in a.items():
            for struct, occ in b.items():
                occ = list(set(occ))
                num = len(occ)
                total_labels[label] += num
                total_structs[struct] += num
                total_s += num
        short = f'{total_s} phrases: ' + ', '.join([f'{v} "{k}"' for k, v in total_labels.items()]) + '.'
        long = '\n'.join([f'{v} — {k}' for k, v in total_structs.items()])
        return short, long

    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.bold = True

    sent_style = styles.add_style('sent', WD_STYLE_TYPE.CHARACTER)
    sent_font = sent_style.font
    sent_font.size = Pt(12)
    sent_font.name = 'Lato'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(11)

    for sent, a in structures.items():
        parts = sent.split(' — sent.')
        doc.add_heading(f'{parts[1]}. {parts[0]}', 0)
        # A. summary
        short, long = gen_summary(a)
        p = doc.add_paragraph()
        p.add_run(short, style=freq_style).add_break()
        p.add_run(long, style=sent_style)

        for phrase, b in a.items():
            for p_struct, c in b.items():
                c = list(set(c))
                for occ in c:
                    _, subtree, label = occ
                    label = label.replace('\n', '')
                    p = doc.add_paragraph()
                    p.add_run(label, style=freq_style)
                    if label.replace('[', '(').replace(']', ')') != subtree:
                        p.add_run().add_break()
                        p.add_run(subtree, style=tree_style)

    doc.save(out_file)


def gen_structure_total_report(structures, out_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.name = 'Lato'
    freq_font.italic = True

    sent_style = styles.add_style('sent', WD_STYLE_TYPE.CHARACTER)
    sent_font = sent_style.font
    sent_font.size = Pt(12)
    sent_font.name = 'Lato'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(11)

    total = order_structures(structures)
    # SUMMARY
    doc.add_heading('Summary', 0)

    total_token = sum([t['freq'] for t in total.values()])

    sorted_structs = sorted([(total[t]['freq'], t) for t in total.keys()], reverse=True)
    rules = []
    total_type = 0
    for n, s in enumerate(sorted_structs):
        main_rule = f'{n+1}.  "{s[1]}": {s[0]}'
        sub_rules = sorted([(total[s[1]]['data'][t]['freq'], t) for t in total[s[1]]['data'].keys()], reverse=True)
        total_type += len(sub_rules)
        sub_rules = '\n'.join([f'\t{n+1})   {s[1]} — {s[0]}' for n, s in enumerate(sub_rules)])
        rules.append((main_rule, sub_rules))

    total_short = f'This text contains {total_token} occurrences of {total_type} different grammatical structures.'
    p = doc.add_paragraph()
    p.add_run(total_short, style=freq_style).add_break()

    for main_rule, sub_rules in rules:
        p.add_run(main_rule, style=sent_style).add_break()
        p.add_run(sub_rules, style=tree_style).add_break()

    # SECTIONS
    for phrase_num, ss in enumerate(sorted_structs):
        freq, phrase = ss
        a = total[phrase]
        doc.add_heading(f'{phrase_num+1}. {phrase}', 0)
        # A. summary
        short = f'"{phrase}" occurs {a["freq"]} times.'
        p = doc.add_paragraph()
        p.add_run(short, style=freq_style).add_break()
        sorted_labels = sorted([(a['data'][t]['freq'], t) for t in a['data'].keys()], reverse=True)
        long = '\n'.join([f'{n + 1})   {l[1]} — {l[0]}' for n, l in enumerate(sorted_labels)])
        p.add_run(long, style=tree_style)

        # B. details
        l_num = 0
        for _, sl in enumerate(sorted_labels):
            freq, label = sl
            l_num += 1
            doc.add_heading(f'{l_num}) {label}', 3)

            p = doc.add_paragraph()
            b = a['data'][label]
            for num, sent in enumerate(b['data']):
                sent = sent.replace('\n', '')
                p.add_run(f'{num+1}.\t').size = Pt(17)
                p.add_run(sent, style=tree_style).add_break()

    doc.save(out_file)


def order_structures(structures):
    total = {}
    for sent, a in structures.items():
        for phrase, b in a.items():
            if phrase not in total:
                total[phrase] = {'freq': 0, 'data': {}}

            for p_struct, c in b.items():
                if p_struct not in total[phrase]['data']:
                    total[phrase]['data'][p_struct] = {'freq': 0, 'data': []}

                for occ in c:
                    location, subtree, _ = occ

                    # add structure to total if not yet there
                    if location not in total[phrase]['data'][p_struct]['data']:
                        total[phrase]['data'][p_struct]['data'].append(location)

                    # increment frequencies
                    total[phrase]['freq'] += 1
                    total[phrase]['data'][p_struct]['freq'] += 1

    return total
