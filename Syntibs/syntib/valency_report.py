import re

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def get_valency_structures(tree, val_structs, total_val_structs):
    def filter_sent(tree):
        return 'sent' in tree.label()

    sents = list(tree.subtrees(filter_sent))
    for sent in sents:
        verb = []
        label_min = []
        label_full = []
        parts = []
        for phrase in sent:
            if '-' in phrase.label():
                l, node = phrase.label().split('-')
                label_full.append(f'{l}({node})')
                if l.isupper():
                    label_min.append(f'{l}({node})')
                parts.append(f'{phrase.label()}({"".join(phrase.leaves())})')
            elif phrase.label() == 'vp':
                for p in phrase.pos():
                    if p[1] == 'VERB' or p[1] == 'NOUN':
                        label_min.append(p[0])
                        label_full.append(p[0])
                        verb.append(p[0])
                parts.append(f'VERB({"".join(phrase.leaves())})')

        label_min = ' '.join(label_min)
        label_full = ' '.join(label_full)
        label = (label_min, label_full)
        verb = ''.join(verb)
        parts = ' '.join(parts)

        if not verb:
            continue

        # create structure if missing
        if verb not in val_structs:
            val_structs[verb] = {}
        if verb not in total_val_structs:
            total_val_structs[verb] = {}

        if label not in val_structs[verb]:
            val_structs[verb][label] = {}
        if label not in total_val_structs[verb]:
            total_val_structs[verb][label] = {}

        if parts not in val_structs[verb][label]:
            val_structs[verb][label][parts] = []
        if parts not in total_val_structs[verb][label]:
            total_val_structs[verb][label][parts] = []

        # increment freq + add parts
        sent_str = ' '.join(sent.leaves())
        if sent_str not in val_structs[verb][label][parts]:
            val_structs[verb][label][parts].append(sent_str)
        if sent_str not in total_val_structs[verb][label][parts]:
            total_val_structs[verb][label][parts].append(sent_str)


def gen_valency_report(val_structs, out_file):
    doc = Document()
    styles = doc.styles

    freq_style = styles.add_style('freq', WD_STYLE_TYPE.CHARACTER)
    freq_font = freq_style.font
    freq_font.italic = True

    sent_style = styles.add_style('sent', WD_STYLE_TYPE.CHARACTER)
    sent_font = sent_style.font
    sent_font.size = Pt(12)
    sent_font.name = 'Lato'

    tree_style = styles.add_style('tree', WD_STYLE_TYPE.CHARACTER)
    tree_font = tree_style.font
    tree_font.name = 'Lato Light'
    tree_font.size = Pt(13)

    count = 1
    for verb, structs in val_structs.items():
        doc.add_heading(f'{count} {verb}', 0)
        count += 1
        n = 1
        for struct, sents in structs.items():
            doc.add_heading(str(n), 2)
            n += 1
            p = doc.add_paragraph('Valency: ')
            p.add_run(struct[0], style=tree_style)
            if struct[0] != struct[1]:
                p = doc.add_paragraph('Extended: ')
                p.add_run(struct[1], style=tree_style)

            doc.add_heading('Occurences', 5)
            for sent, sents_str in sents.items():
                p = doc.add_paragraph()
                p.add_run(sent + ': ', style=sent_style)
                strs = '"' + '", "'.join(sents_str) + '"'
                p.add_run(strs, style=tree_style)

    doc.save(out_file)
